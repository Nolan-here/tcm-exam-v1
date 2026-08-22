#!/usr/bin/env python3
"""将 2024 年一试 Word 题库转换为浏览器可直接加载的 JavaScript 数据。

Word 的自动编号不属于段落正文，因此这里同时解析 document.xml 与
numbering.xml。脚本只读取源文档，不会修改它。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


UNIT_RE = re.compile(r"^第([一二三四])单元$")
ANSWER_RE = re.compile(r"答案\s*[：:]\s*([A-E])")
EXPLANATION_RE = re.compile(r"解析\s*[：:]\s*(.*)", re.S)
QUESTION_RE = re.compile(r"^\s*(\d{1,3})[.．、]\s*(.*)$", re.S)
OPTION_RE = re.compile(r"(?:^|\n|\s)([A-E])[.．、]\s*")
UNIT_MAP = {"一": 1, "二": 2, "三": 3, "四": 4}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


class NumberingReader:
    def __init__(self, document: Document):
        root = document.part.numbering_part.element
        self.abstracts = {
            node.get(qn("w:abstractNumId")): node
            for node in root.findall(qn("w:abstractNum"))
        }
        self.nums = {}
        for node in root.findall(qn("w:num")):
            num_id = node.get(qn("w:numId"))
            abstract_node = node.find(qn("w:abstractNumId"))
            if abstract_node is None:
                continue
            overrides = {}
            for override in node.findall(qn("w:lvlOverride")):
                level = int(override.get(qn("w:ilvl"), "0"))
                start_node = override.find(qn("w:startOverride"))
                if start_node is not None:
                    overrides[level] = int(start_node.get(qn("w:val"), "1"))
            self.nums[num_id] = (abstract_node.get(qn("w:val")), overrides)
        self.counters: dict[tuple[str, int], int] = {}

    def label_for(self, paragraph) -> str | None:
        ppr = paragraph._p.pPr
        num_pr = None if ppr is None else ppr.numPr
        if num_pr is None or num_pr.numId is None:
            return None
        num_id = str(num_pr.numId.val)
        level = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
        if num_id not in self.nums:
            return None
        abstract_id, overrides = self.nums[num_id]
        abstract = self.abstracts.get(abstract_id)
        if abstract is None:
            return None
        level_node = next(
            (
                node
                for node in abstract.findall(qn("w:lvl"))
                if int(node.get(qn("w:ilvl"), "0")) == level
            ),
            None,
        )
        if level_node is None:
            return None
        fmt_node = level_node.find(qn("w:numFmt"))
        start_node = level_node.find(qn("w:start"))
        fmt = "decimal" if fmt_node is None else fmt_node.get(qn("w:val"), "decimal")
        start = overrides.get(
            level,
            1 if start_node is None else int(start_node.get(qn("w:val"), "1")),
        )
        key = (num_id, level)
        value = self.counters.get(key, start - 1) + 1
        self.counters[key] = value
        for deeper_key in [key for key in self.counters if key[0] == num_id and key[1] > level]:
            del self.counters[deeper_key]
        if fmt == "decimal":
            return f"{value}."
        if fmt == "upperLetter" and 1 <= value <= 26:
            return f"{chr(64 + value)}."
        return None


def split_embedded_questions(text: str) -> list[str]:
    """拆分少数粘在上一题解析末尾的手工题号。"""
    matches = list(re.finditer(r"(?<!\d)(\d{1,3})[．、](?=\s*[^\d])", text))
    if not matches:
        return [text]
    starts = []
    for match in matches:
        # 文档里的全角点只用于手工题号；药物剂量和小数均使用半角点。
        starts.append(match.start())
    if not starts:
        return [text]
    result = []
    cursor = 0
    for position in starts:
        if position > cursor:
            result.append(text[cursor:position].strip())
        cursor = position
    result.append(text[cursor:].strip())
    return [part for part in result if part]


def parse_options(text: str) -> dict[str, str]:
    matches = list(OPTION_RE.finditer(text))
    options = {}
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        value = text[match.end() : end]
        value = re.split(r"答案\s*[：:]|解析\s*[：:]", value, maxsplit=1)[0]
        options[match.group(1)] = clean(value)
    return options


def paragraph_events(document: Document):
    numbering = NumberingReader(document)
    for index, paragraph in enumerate(document.paragraphs):
        text = clean(paragraph.text)
        label = numbering.label_for(paragraph)
        if not text and not label:
            continue
        combined = clean(f"{label or ''} {text}")
        # 原文两处 E 选项被错误套用了十进制自动编号，按正文语义还原。
        if re.fullmatch(r"1\. 岁(?:半)?答案[：:][A-E]", combined):
            combined = f"E. {combined.removeprefix('1. ')}"
        for part in split_embedded_questions(combined):
            yield index, part


def extract_questions(document: Document):
    unit = 0
    question_type = ""
    shared_stem = ""
    shared_options: dict[str, str] = {}
    pending_shared_options = False
    shared_stem_group = ""
    shared_options_group = ""
    group_serials = Counter()
    current = None
    questions = []

    def finish():
        nonlocal current
        if current is None:
            return
        body = "\n".join(current["parts"])
        answer_match = ANSWER_RE.search(body)
        explanation_match = EXPLANATION_RE.search(body)
        options = parse_options(body)
        if len(options) < 5 and current["shared_options"]:
            options = dict(current["shared_options"])
        first_option = OPTION_RE.search(body)
        stem_end = first_option.start() if first_option else len(body)
        stem = re.split(r"答案\s*[：:]|解析\s*[：:]", body[:stem_end], maxsplit=1)[0]
        prompt = clean(stem)
        question = {
            "id": f"2024-U{current['unit']}-{current['number']:03d}",
            "unit": current["unit"],
            "number": current["number"],
            "type": current["type"],
            "stem": clean(f"{current['shared_stem']} {prompt}"),
            "options": options,
            "answer": answer_match.group(1) if answer_match else "",
            "explanation": clean(explanation_match.group(1)) if explanation_match else "",
            "sourceParagraph": current["paragraph"],
        }
        group_id = current["shared_stem_group"] or current["shared_options_group"]
        if group_id:
            question.update({"prompt": prompt, "groupId": group_id})
            if current["shared_stem_group"]:
                question["sharedStem"] = current["shared_stem"]
            else:
                question["sharedOptions"] = dict(current["shared_options"])
        questions.append(question)
        current = None

    for paragraph_index, text in paragraph_events(document):
        unit_match = UNIT_RE.match(text)
        if unit_match:
            finish()
            unit = UNIT_MAP[unit_match.group(1)]
            question_type = ""
            shared_stem = ""
            shared_options = {}
            pending_shared_options = False
            shared_stem_group = ""
            shared_options_group = ""
            continue
        heading_match = re.fullmatch(r"(A1/A2|A3/A4|B)\s*型题", text)
        if heading_match:
            finish()
            question_type = {"A3/A4": "A3", "B": "B1"}.get(
                heading_match.group(1), heading_match.group(1)
            )
            shared_stem = ""
            shared_options = {}
            pending_shared_options = False
            shared_stem_group = ""
            shared_options_group = ""
            continue
        if text.startswith("共用题干："):
            finish()
            shared_stem = clean(text.removeprefix("共用题干："))
            group_serials[(unit, "A3")] += 1
            shared_stem_group = f"2024-U{unit}-A3-G{group_serials[(unit, 'A3')]:03d}"
            continue
        if match := re.search(r"（(\d+)\s*[～~-]\s*(\d+)\s*题共用备选答案）", text):
            finish()
            shared_options = parse_options(text)
            pending_shared_options = len(shared_options) < 5
            shared_options_group = (
                f"2024-U{unit}-B1-{int(match.group(1)):03d}-{int(match.group(2)):03d}"
            )
            continue

        question_match = QUESTION_RE.match(text)
        if question_match and unit:
            finish()
            current = {
                "unit": unit,
                "number": int(question_match.group(1)),
                "type": question_type,
                "shared_stem": shared_stem,
                "shared_stem_group": shared_stem_group,
                "shared_options": dict(shared_options),
                "shared_options_group": shared_options_group,
                "paragraph": paragraph_index,
                "parts": [question_match.group(2)],
            }
            continue

        if pending_shared_options:
            candidate = parse_options(text)
            if candidate:
                shared_options.update(candidate)
                if len(shared_options) == 5:
                    pending_shared_options = False
                continue
        if current is not None:
            current["parts"].append(text)

    finish()
    apply_known_source_repairs(questions, document)
    for question in questions:
        question["options"] = {
            letter: question["options"].get(letter, "") for letter in "ABCDE"
        }
    questions.sort(key=lambda question: (question["unit"], question["number"]))
    groups = {}
    for question in questions:
        if question.get("groupId"):
            groups.setdefault(question["groupId"], []).append(question)
    for members in groups.values():
        start = min(question["number"] for question in members)
        end = max(question["number"] for question in members)
        for question in members:
            question["groupStart"] = start
            question["groupEnd"] = end
    return questions


def apply_known_source_repairs(questions: list[dict], document: Document):
    """修复原文中段落顺序交叉或编号样式中断的已知位置。

    这些修复均以 Word 源段落号为锚点，并由后续完整性测试约束，避免
    静默猜测缺失内容。
    """
    by_id = {question["id"]: question for question in questions}
    paragraphs = [clean(paragraph.text) for paragraph in document.paragraphs]

    def joined(*indexes: int) -> str:
        return clean(" ".join(paragraphs[index] for index in indexes))

    def set_options(question_id: str, *indexes: int):
        by_id[question_id]["options"] = parse_options("\n".join(paragraphs[index] for index in indexes))

    # 第二单元第 49～51 题的三段解析在源文档中交叉排列。
    by_id["2024-U2-049"]["explanation"] = clean(
        f"{by_id['2024-U2-049']['explanation']} {paragraphs[721]}"
    )
    by_id["2024-U2-050"]["explanation"] = by_id["2024-U2-050"]["explanation"].split(
        " 解析：", 1
    )[0]
    by_id["2024-U2-051"]["explanation"] = clean(paragraphs[732].removeprefix("解析："))

    # 第三单元第 71、73 题的选项和答案被 A3/A4 标题分隔。
    set_options("2024-U3-071", 1118, 1128)
    by_id["2024-U3-071"]["options"]["A"] = clean(paragraphs[1118])
    by_id["2024-U3-071"]["answer"] = "A"
    by_id["2024-U3-071"]["explanation"] = clean(paragraphs[1129].removeprefix("解析："))
    set_options("2024-U3-073", 1125, 1126, 1127)
    by_id["2024-U3-073"]["answer"] = "D"
    by_id["2024-U3-073"]["explanation"] = clean(paragraphs[1134].removeprefix("解析："))

    # 第三单元第 86 题的后三个选项和解析被 B 型题标题分隔。
    source_86 = re.sub(r"^A[.．、]\s*", "", paragraphs[1192])
    set_options("2024-U3-086", 1204)
    option_b = re.search(r"\sB[.．、]\s*", source_86)
    if option_b:
        by_id["2024-U3-086"]["options"]["A"] = clean(source_86[: option_b.start()])
        by_id["2024-U3-086"]["options"]["B"] = clean(source_86[option_b.end() :])
    by_id["2024-U3-086"]["answer"] = "A"
    by_id["2024-U3-086"]["explanation"] = clean(paragraphs[1205].removeprefix("解析："))

    # 第 89 题解析被第 86 题的补排段落截断。
    by_id["2024-U3-089"]["explanation"] = clean(
        f"{paragraphs[1203].removeprefix('解析：')} {paragraphs[1209]}"
    )

    # 第四单元第 59 题被误插入的 B 型题标题从中间截断。
    source_59_a = paragraphs[1550].split("A．", 1)[-1]
    source_59_e = re.sub(r"^E[.．、]\s*", "", paragraphs[1554])
    source_59_e = re.split(r"答案\s*[：:]", source_59_e, maxsplit=1)[0]
    middle_59 = parse_options(paragraphs[1553])
    by_id["2024-U4-059"]["options"] = {
        "A": clean(source_59_a),
        "B": clean(re.sub(r"^B[.．、]\s*", "", paragraphs[1551])),
        "C": middle_59["C"],
        "D": middle_59["D"],
        "E": clean(source_59_e),
    }
    by_id["2024-U4-059"]["answer"] = "D"
    by_id["2024-U4-059"]["explanation"] = clean(paragraphs[1555].removeprefix("解析："))

    # 第 59～61 题实际共用同一个闭经病例题干；误插的 B 型题标题曾使
    # 第 60、61 题被错误分型。恢复为完整 A3 题组。
    shared_59 = by_id["2024-U4-059"]["sharedStem"]
    group_59 = by_id["2024-U4-059"]["groupId"]
    for question_id in ("2024-U4-060", "2024-U4-061"):
        question = by_id[question_id]
        question["type"] = "A3"
        question["prompt"] = question["stem"]
        question["stem"] = clean(f"{shared_59} {question['prompt']}")
        question["groupId"] = group_59
        question["sharedStem"] = shared_59


def validate(questions: list[dict]) -> list[str]:
    errors = []
    ids = Counter(question["id"] for question in questions)
    for duplicate, count in ids.items():
        if count > 1:
            errors.append(f"重复题号：{duplicate}（{count} 次）")
    for question in questions:
        missing = [key for key in "ABCDE" if not question["options"].get(key)]
        if not question["stem"]:
            errors.append(f"{question['id']} 缺少题干（源段落 {question['sourceParagraph']}）")
        if missing:
            errors.append(
                f"{question['id']} 缺少选项 {','.join(missing)}（源段落 {question['sourceParagraph']}）"
            )
        if question["answer"] not in "ABCDE":
            errors.append(f"{question['id']} 缺少有效答案（源段落 {question['sourceParagraph']}）")
        if not question["explanation"]:
            errors.append(f"{question['id']} 缺少解析（源段落 {question['sourceParagraph']}）")
    expected = {1: 136, 2: 68, 3: 108, 4: 79}
    actual = Counter(question["unit"] for question in questions)
    for unit, count in expected.items():
        if actual[unit] != count:
            errors.append(f"第 {unit} 单元题数应为 {count}，实际为 {actual[unit]}")
    return errors


def write_javascript(path: Path, questions: list[dict], source_name: str):
    payload = json.dumps(questions, ensure_ascii=False, separators=(",", ":"))
    unit_counts = Counter(question["unit"] for question in questions)
    text = (
        "// 此文件由 scripts/import_2024_docx.py 生成，请勿手工编辑。\n"
        f"export const QUESTION_BANK_VERSION = '2024-docx-v2';\n"
        f"export const QUESTION_BANK_SOURCE = {json.dumps(source_name, ensure_ascii=False)};\n"
        f"export const QUESTIONS_2024 = {payload};\n"
        "export const EXAM_UNITS = "
        + json.dumps(
            [
                {"unit": unit, "name": f"第{UNIT_NAMES[unit]}单元", "count": unit_counts[unit]}
                for unit in range(1, 5)
            ],
            ensure_ascii=False,
            separators=(",", ":"),
        )
        + ";\n"
        "const QUESTION_BY_ID = new Map(QUESTIONS_2024.map((question) => [question.id, question]));\n"
        "export function getQuestionById(id) { return QUESTION_BY_ID.get(id) || null; }\n"
        "export function getQuestionsForUnit(unit) { return QUESTIONS_2024.filter((question) => question.unit === unit); }\n"
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8", newline="\n")


UNIT_NAMES = {1: "一", 2: "二", 3: "三", 4: "四"}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--report-only", action="store_true")
    args = parser.parse_args()
    if not args.source.is_file():
        print(f"找不到源文档：{args.source}", file=sys.stderr)
        return 2
    questions = extract_questions(Document(args.source))
    errors = validate(questions)
    counts = Counter(question["unit"] for question in questions)
    print(f"共解析 {len(questions)} 题；单元分布：{dict(sorted(counts.items()))}")
    if errors:
        print(f"发现 {len(errors)} 个完整性问题：", file=sys.stderr)
        for error in errors:
            print(f"- {error}", file=sys.stderr)
        return 1
    if not args.report_only:
        write_javascript(args.output, questions, args.source.name)
        print(f"已生成：{args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
